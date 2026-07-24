"""Convert a competitive-intel brief (Markdown) to a Word document.

Handles the subset of Markdown these briefs actually use: headers (#-###),
bullet lists, pipe tables, bold spans, and plain paragraphs. Not a general
Markdown renderer -- this repo's brief format is deliberately constrained.
"""
import pathlib
import re
import sys

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_bold_runs(paragraph, text):
    for chunk in re.split(r"(\*\*.+?\*\*)", text):
        if chunk.startswith("**") and chunk.endswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            run.bold = True
        elif chunk:
            paragraph.add_run(chunk)


def is_table_row(line: str) -> bool:
    return line.strip().startswith("|") and line.strip().endswith("|")


def is_table_sep(line: str) -> bool:
    return bool(re.match(r"^\s*\|?[\s:|-]+\|?\s*$", line)) and "-" in line


def convert(md_path: pathlib.Path, docx_path: pathlib.Path, title: str):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    lines = md_path.read_text().splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph("_" * 40)
            i += 1
            continue

        m = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if m:
            level = min(len(m.group(1)) + 1, 4)  # doc title already used level 0
            doc.add_heading(m.group(2), level=level)
            i += 1
            continue

        if is_table_row(stripped):
            table_lines = []
            while i < len(lines) and is_table_row(lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1
            rows = [
                [c.strip() for c in row.strip("|").split("|")]
                for row in table_lines
                if not is_table_sep(row)
            ]
            if rows:
                ncols = len(rows[0])
                table = doc.add_table(rows=0, cols=ncols)
                table.style = "Light Grid Accent 1"
                for r_idx, row in enumerate(rows):
                    cells = table.add_row().cells
                    for c_idx in range(ncols):
                        text = row[c_idx] if c_idx < len(row) else ""
                        p = cells[c_idx].paragraphs[0]
                        add_bold_runs(p, text)
                        if r_idx == 0:
                            for run in p.runs:
                                run.bold = True
            continue

        bullet_m = re.match(r"^(\s*)-\s+(.*)", line)
        if bullet_m:
            indent = len(bullet_m.group(1))
            style_name = "List Bullet2" if indent >= 2 else "List Bullet"
            p = doc.add_paragraph(style=style_name)
            add_bold_runs(p, bullet_m.group(2))
            i += 1
            continue

        num_m = re.match(r"^\s*\d+\.\s+(.*)", line)
        if num_m:
            p = doc.add_paragraph(style="List Number")
            add_bold_runs(p, num_m.group(1))
            i += 1
            continue

        p = doc.add_paragraph()
        add_bold_runs(p, stripped)
        i += 1

    doc.save(docx_path)
    print(f"[write] {docx_path}")


if __name__ == "__main__":
    root = pathlib.Path(__file__).parent.parent
    jobs = [
        (root / "briefs" / "gcp-earnings-2026-07-22.md",
         root / "briefs" / "gcp-earnings-2026-07-22.docx",
         "Google Cloud Earnings Brief — Q2 2026 (AWS Field Team)"),
        (root / "briefs" / "gcp-technical-pricing-2026-07-22.md",
         root / "briefs" / "gcp-technical-pricing-2026-07-22.docx",
         "Google Cloud Technical & Pricing Battlecards — 2026-07-22"),
    ]
    for md, out, title in jobs:
        convert(md, out, title)
